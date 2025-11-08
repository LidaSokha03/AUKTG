from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from app.db.models.cv import CV


class DOCXTemplate:
    """Base class for DOCX templates"""
    
    def __init__(self, cv_obj: CV):
        self.cv = cv_obj
        self.doc = Document()
    
    def generate(self, output_path: Path) -> Path:
        raise NotImplementedError


class ClassicDOCXTemplate(DOCXTemplate):
    """Classic DOCX template - your current one"""
    
    def generate(self, output_path: Path) -> Path:
        doc = self.doc
        
        # Name heading
        name_heading = doc.add_heading(f"{self.cv.firstname} {self.cv.lastname}", level=0)
        name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.add_run(f"Email: {self.cv.email}  |  ").bold = False
        contact_p.add_run(f"{self.cv.phone}").bold = False
        
        doc.add_paragraph("")
        
        # Experience
        doc.add_heading("Experience", level=1)
        doc.add_paragraph(self.cv.experience, style="List Paragraph")
        doc.add_paragraph("")
        
        # Education
        doc.add_heading("Education", level=1)
        doc.add_paragraph(self.cv.education, style="List Paragraph")
        doc.add_paragraph("")
        
        # Courses and Languages
        doc.add_heading("Courses and Languages", level=1)
        doc.add_paragraph(self.cv.courses, style="List Paragraph")
        doc.add_paragraph("")
        
        # Skills
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(self.cv.skills, style="List Paragraph")
        doc.add_paragraph("")
        
        doc.save(output_path)
        return output_path


class ModernDOCXTemplate(DOCXTemplate):
    """Modern minimalist DOCX template"""
    
    def generate(self, output_path: Path) -> Path:
        doc = self.doc
        
        # Name - large and bold
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(f"{self.cv.firstname} {self.cv.lastname}")
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(41, 128, 185)  # Blue
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_p = doc.add_paragraph()
        contact_run = contact_p.add_run(f"{self.cv.email} | {self.cv.phone}")
        contact_run.font.size = Pt(11)
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        
        # Sections with modern styling
        sections = [
            ("EXPERIENCE", self.cv.experience),
            ("EDUCATION", self.cv.education),
            ("COURSES & LANGUAGES", self.cv.courses),
            ("SKILLS", self.cv.skills)
        ]
        
        for title, content in sections:
            # Section heading
            heading_p = doc.add_paragraph()
            heading_run = heading_p.add_run(title)
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(41, 128, 185)
            
            # Content
            content_p = doc.add_paragraph(content)
            content_p.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph("")
        
        doc.save(output_path)
        return output_path


class ProfessionalDOCXTemplate(DOCXTemplate):
    """Professional business DOCX template"""
    
    def generate(self, output_path: Path) -> Path:
        doc = self.doc
        
        # Header section with line
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(f"{self.cv.firstname} {self.cv.lastname}")
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Contact info
        contact_p = doc.add_paragraph()
        contact_run = contact_p.add_run(f"{self.cv.email} | {self.cv.phone}")
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Add horizontal line
        doc.add_paragraph("_" * 80)
        
        # Sections
        sections = [
            ("Professional Experience", self.cv.experience),
            ("Education", self.cv.education),
            ("Skills & Languages", self.cv.skills),
            ("Courses & Certifications", self.cv.courses)
        ]
        
        for title, content in sections:
            if content:  # Only add if content exists
                # Heading
                heading = doc.add_heading(title, level=1)
                heading_format = heading.runs[0]
                heading_format.font.size = Pt(13)
                heading_format.font.color.rgb = RGBColor(0, 0, 0)
                
                # Content
                doc.add_paragraph(content)
                doc.add_paragraph("")
        
        doc.save(output_path)
        return output_path


# Template registry
DOCX_TEMPLATES = {
    'classic': ClassicDOCXTemplate,
    'modern': ModernDOCXTemplate,
    'professional': ProfessionalDOCXTemplate
}


def cv_to_docx_template(cv_obj: CV, template: str = 'classic', out_path: Path = Path('exports')) -> Path:
    """Generate CV DOCX with selected template"""
    out_path.mkdir(exist_ok=True)
    file_path = out_path / f"cv_{cv_obj.user_id}_{template}.docx"
    
    template_class = DOCX_TEMPLATES.get(template, ClassicDOCXTemplate)
    template_instance = template_class(cv_obj)
    
    return template_instance.generate(file_path)