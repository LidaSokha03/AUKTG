from docx import Document
from pathlib import Path
from app.db.models.cv import CV


def cv_to_docx(cv_obj: CV, out_path: Path = Path("exports")) -> Path:
    """
    Генерує CV у форматі DOCX (Microsoft Word)
    """
    out_path.mkdir(exist_ok=True)
    file_path = out_path / f"cv_{cv_obj.user_id}.docx"

    doc = Document()

    name_heading = doc.add_heading(f"{cv_obj.firstname} {cv_obj.lastname}", level=0)
    name_heading.alignment = 1

    contact_p = doc.add_paragraph()
    contact_p.alignment = 1
    contact_p.add_run(f"Email: {cv_obj.email}  |  ").bold = False
    contact_p.add_run(f"{cv_obj.phone}").bold = False

    doc.add_paragraph("")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph(cv_obj.experience, style="List Paragraph")
    doc.add_paragraph("")

    doc.add_heading("Education", level=1)
    doc.add_paragraph(cv_obj.education, style="List Paragraph")
    doc.add_paragraph("")

    doc.add_heading("Courses and Languages", level=1)
    doc.add_paragraph(cv_obj.courses, style="List Paragraph")
    doc.add_paragraph("")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(cv_obj.skills, style="List Paragraph")
    doc.add_paragraph("")

    doc.save(file_path)
    return file_path
